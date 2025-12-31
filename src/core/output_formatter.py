from __future__ import annotations

import json
import logging
import os
from typing import Any


def save_output(text: str, filename: str, style: str = "default") -> None:
    Document: Any = None
    try:
        from docx import Document
    except ImportError:
        pass

    base, ext = os.path.splitext(filename)

    if ext == ".txt":
        with open(filename, "w", encoding="utf-8") as f:
            f.write(text)

    elif ext == ".docx":
        if Document is None:
            raise ImportError("A biblioteca 'python-docx' é necessária para salvar em .docx. Instale com 'pip install python-docx'")
        doc = Document()
        if style == "bullets":
            for line in text.split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip(), style="List Bullet")
        else:
            doc.add_paragraph(text)
        doc.save(filename)

    elif ext == ".md":
        with open(filename, "w", encoding="utf-8") as f:
            if style == "bullets":
                for line in text.split("\n"):
                    if line.strip():
                        f.write(f"- {line.strip()}\n")
            else:
                f.write(text)

    elif ext == ".json":
        output_data = {"style": style, "content": text}
        with open(filename, "w", encoding="utf-8") as f:
            json.dump(output_data, f, ensure_ascii=False, indent=4)

    else:
        logging.warning(f"Formato '{ext}' não suportado diretamente. Salvando como texto puro.")
        with open(filename, "w", encoding="utf-8") as f:
            f.write(text)

    logging.info(f"Arquivo salvo em '{filename}' com o estilo '{style}'.")


# "A pena e mais poderosa que a espada." - Edward Bulwer-Lytton
