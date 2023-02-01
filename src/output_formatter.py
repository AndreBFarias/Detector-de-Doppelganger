# 56
import os
import json
import logging

def save_output(text, filename, style="default"):
    """
    Salva o texto no formato especificado, aplicando um estilo.
    """
    try:
        # 56
        # Tenta importar docx apenas quando necessário
        from docx import Document
    except ImportError:
        Document = None

    # 56
    base, ext = os.path.splitext(filename)
    
    if ext == ".txt":
        # 57
        with open(filename, "w", encoding="utf-8") as f:
            f.write(text)

    elif ext == ".docx":
        # 58
        if not Document:
            raise ImportError("A biblioteca 'python-docx' é necessária para salvar em .docx. Instale com 'pip install python-docx'")
        doc = Document()
        if style == "bullets":
            for line in text.split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip(), style="List Bullet")
        else:
            doc.add_paragraph(text)
        doc.save(filename)

    elif ext == ".md":
        # 59
        with open(filename, "w", encoding="utf-8") as f:
            if style == "bullets":
                for line in text.split("\n"):
                    if line.strip():
                        f.write(f"- {line.strip()}\n")
            else:
                f.write(text)

    elif ext == ".json":
        # 60
        output_data = {"style": style, "content": text}
        with open(filename, "w", encoding="utf-8") as f:
            json.dump(output_data, f, ensure_ascii=False, indent=4)

    else:
        # 60
        # Fallback para qualquer outra extensão é salvar como texto puro
        logging.warning(f"Formato '{ext}' não suportado diretamente. Salvando como texto puro.")
        with open(filename, "w", encoding="utf-8") as f:
            f.write(text)
            
    # 60
    logging.info(f"Arquivo salvo em '{filename}' com o estilo '{style}'.")
